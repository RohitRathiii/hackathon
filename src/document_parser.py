#!/usr/bin/env python3
"""
Document Parser Module for LLM Document Query System

This module implements structure-aware document parsing for PDF, DOCX, and email formats,
following the methodology that discourages fixed-size chunking and recommends 
structure-aware methods like Markdown splitting for legal documents.

Key Features:
- PDF parsing via PyMuPDF with structure preservation
- DOCX parsing via python-docx for sections and tables  
- Email parsing for .eml format
- URL download capability
- Semantic chunking avoiding mid-sentence splits
- Explainability through detailed logging
"""

import os
import re
import email
import logging
from pathlib import Path
from typing import List, Dict, Any, Optional, Union
from dataclasses import dataclass
from urllib.parse import urlparse
import tempfile

# Document processing imports
import fitz  # PyMuPDF - Using PyMuPDF as recommended for superior PDF handling [doc ref: Document Parsing section]
import docx  # python-docx for DOCX document parsing
import requests  # HTTP library for document URL downloads
import nltk  # Natural language processing for semantic refinement
import spacy  # Advanced NLP for structure-aware chunking

# Configure logging for parsing explainability
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

@dataclass
class ChunkMetadata:
    """Metadata for document chunks with rich contextual information"""
    document_id: str
    page_number: Optional[int] = None
    section: Optional[str] = None
    chunk_type: str = "paragraph"  # paragraph, clause, table, header, etc.
    confidence: float = 1.0
    source_file: Optional[str] = None
    extraction_method: Optional[str] = None

@dataclass
class DocumentChunk:
    """Structured document chunk with content and metadata"""
    content: str
    metadata: ChunkMetadata
    
    def __post_init__(self):
        """Validate chunk content and metadata"""
        if not self.content.strip():
            raise ValueError("Chunk content cannot be empty")
        if len(self.content) > 10000:  # Memory efficiency constraint
            logger.warning(f"Large chunk detected: {len(self.content)} characters")

class DocumentParser:
    """
    Main document parser orchestrator that handles multiple document formats
    with structure preservation and semantic chunking
    """
    
    def __init__(self):
        """Initialize parser with NLP models for semantic refinement"""
        self.supported_formats = {'.pdf', '.docx', '.eml', '.txt'}
        self._setup_nlp_models()
        logger.info("DocumentParser initialized with structure-aware chunking")
    
    def _setup_nlp_models(self):
        """Setup NLTK and spaCy models for semantic analysis"""
        try:
            # Download required NLTK data
            nltk.download('punkt', quiet=True)
            nltk.download('stopwords', quiet=True)
            
            # Load spaCy model for semantic refinement
            try:
                self.nlp = spacy.load("en_core_web_sm")
            except OSError:
                logger.warning("spaCy model 'en_core_web_sm' not found. Using basic tokenization.")
                self.nlp = None
                
        except Exception as e:
            logger.error(f"Error setting up NLP models: {e}")
            self.nlp = None
    
    def parse_document(self, file_path_or_url: str) -> List[DocumentChunk]:
        """
        Main parsing function that handles file path or URL input
        
        Args:
            file_path_or_url: Local file path or URL to document
            
        Returns:
            List of DocumentChunk objects with structured content and metadata
        """
        logger.info(f"Starting document parsing for: {file_path_or_url}")
        
        # Handle URL downloads
        if self._is_url(file_path_or_url):
            file_path = self.download_document(file_path_or_url)
        else:
            file_path = file_path_or_url
        
        # Detect file type and parse accordingly
        file_type = self.detect_file_type(file_path)
        document_id = self._generate_document_id(file_path)
        
        try:
            if file_type == 'pdf':
                chunks = self._parse_pdf(file_path, document_id)
            elif file_type == 'docx':
                chunks = self._parse_docx(file_path, document_id)
            elif file_type == 'eml':
                chunks = self._parse_email(file_path, document_id)
            elif file_type == 'txt':
                chunks = self._parse_text(file_path, document_id)
            else:
                raise ValueError(f"Unsupported file format: {file_type}")
            
            # Apply semantic refinement to avoid mid-sentence splits
            refined_chunks = self._refine_chunks_semantically(chunks)
            
            logger.info(f"Parsed {len(refined_chunks)} chunks from {file_path}, "
                       f"preserving {file_type.upper()} structure [doc ref: Avoid Fixed-Size Chunking]")
            
            return refined_chunks
            
        except Exception as e:
            logger.error(f"Error parsing document {file_path}: {e}")
            raise
        finally:
            # Clean up downloaded files
            if self._is_url(file_path_or_url) and os.path.exists(file_path):
                os.unlink(file_path)
    
    def download_document(self, url: str) -> str:
        """
        Download document from URL with error handling
        
        Args:
            url: URL to download document from
            
        Returns:
            Path to downloaded temporary file
        """
        logger.info(f"Downloading document from URL: {url}")
        
        try:
            response = requests.get(url, timeout=30, stream=True)
            response.raise_for_status()
            
            # Create temporary file with appropriate extension
            parsed_url = urlparse(url)
            file_extension = Path(parsed_url.path).suffix or '.tmp'
            
            with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as tmp_file:
                for chunk in response.iter_content(chunk_size=8192):
                    tmp_file.write(chunk)
                temp_path = tmp_file.name
            
            logger.info(f"Document downloaded successfully to: {temp_path}")
            return temp_path
            
        except requests.RequestException as e:
            logger.error(f"Failed to download document from {url}: {e}")
            raise
    
    def detect_file_type(self, file_path: str) -> str:
        """
        Detect document format based on file extension and content
        
        Args:
            file_path: Path to the file
            
        Returns:
            File type string (pdf, docx, eml, txt)
        """
        file_extension = Path(file_path).suffix.lower()
        
        type_mapping = {
            '.pdf': 'pdf',
            '.docx': 'docx', 
            '.eml': 'eml',
            '.txt': 'txt',
            '.doc': 'docx'  # Treat .doc as docx for simplicity
        }
        
        detected_type = type_mapping.get(file_extension, 'txt')
        logger.info(f"Detected file type: {detected_type} for {file_path}")
        
        return detected_type
    
    def _parse_pdf(self, file_path: str, document_id: str) -> List[DocumentChunk]:
        """
        Parse PDF using PyMuPDF with structure preservation
        
        Args:
            file_path: Path to PDF file
            document_id: Unique document identifier
            
        Returns:
            List of DocumentChunk objects
        """
        logger.info(f"Parsing PDF with PyMuPDF: {file_path}")
        chunks = []
        
        try:
            doc = fitz.open(file_path)
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                
                # Extract text with layout information
                text_dict = page.get_text("dict")
                
                # Process blocks to preserve structure
                for block in text_dict["blocks"]:
                    if "lines" in block:  # Text block
                        block_text = ""
                        for line in block["lines"]:
                            for span in line["spans"]:
                                block_text += span["text"] + " "
                        
                        if block_text.strip():
                            # Determine chunk type based on formatting
                            chunk_type = self._classify_pdf_block(block, block_text)
                            
                            metadata = ChunkMetadata(
                                document_id=document_id,
                                page_number=page_num + 1,
                                chunk_type=chunk_type,
                                source_file=file_path,
                                extraction_method="pymupdf_structure_aware"
                            )
                            
                            chunks.append(DocumentChunk(
                                content=block_text.strip(),
                                metadata=metadata
                            ))
            
            doc.close()
            logger.info(f"Extracted {len(chunks)} structured blocks from PDF")
            
        except Exception as e:
            logger.error(f"Error parsing PDF {file_path}: {e}")
            raise
        
        return chunks
    
    def _parse_docx(self, file_path: str, document_id: str) -> List[DocumentChunk]:
        """
        Parse DOCX using python-docx for sections and tables
        
        Args:
            file_path: Path to DOCX file
            document_id: Unique document identifier
            
        Returns:
            List of DocumentChunk objects
        """
        logger.info(f"Parsing DOCX with python-docx: {file_path}")
        chunks = []
        
        try:
            doc = docx.Document(file_path)
            
            # Parse paragraphs with style information
            for para_idx, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    # Determine chunk type based on style
                    chunk_type = self._classify_docx_paragraph(paragraph)
                    
                    metadata = ChunkMetadata(
                        document_id=document_id,
                        section=paragraph.style.name if paragraph.style else None,
                        chunk_type=chunk_type,
                        source_file=file_path,
                        extraction_method="python_docx_structure_aware"
                    )
                    
                    chunks.append(DocumentChunk(
                        content=paragraph.text.strip(),
                        metadata=metadata
                    ))
            
            # Parse tables separately
            for table_idx, table in enumerate(doc.tables):
                table_content = self._extract_table_content(table)
                if table_content:
                    metadata = ChunkMetadata(
                        document_id=document_id,
                        section=f"table_{table_idx + 1}",
                        chunk_type="table",
                        source_file=file_path,
                        extraction_method="python_docx_table"
                    )
                    
                    chunks.append(DocumentChunk(
                        content=table_content,
                        metadata=metadata
                    ))
            
            logger.info(f"Extracted {len(chunks)} structured elements from DOCX")
            
        except Exception as e:
            logger.error(f"Error parsing DOCX {file_path}: {e}")
            raise
        
        return chunks
    
    def _parse_email(self, file_path: str, document_id: str) -> List[DocumentChunk]:
        """
        Parse email (.eml) format to extract body and attachments
        
        Args:
            file_path: Path to EML file
            document_id: Unique document identifier
            
        Returns:
            List of DocumentChunk objects
        """
        logger.info(f"Parsing email (.eml): {file_path}")
        chunks = []
        
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                msg = email.message_from_file(f)
            
            # Extract email headers
            headers = {
                'subject': msg.get('Subject', ''),
                'from': msg.get('From', ''),
                'to': msg.get('To', ''),
                'date': msg.get('Date', '')
            }
            
            # Create header chunk
            header_content = f"Subject: {headers['subject']}\nFrom: {headers['from']}\nTo: {headers['to']}\nDate: {headers['date']}"
            
            metadata = ChunkMetadata(
                document_id=document_id,
                section="email_header",
                chunk_type="header",
                source_file=file_path,
                extraction_method="email_parser"
            )
            
            chunks.append(DocumentChunk(
                content=header_content,
                metadata=metadata
            ))
            
            # Extract email body
            if msg.is_multipart():
                for part in msg.walk():
                    if part.get_content_type() == "text/plain":
                        body = part.get_payload(decode=True)
                        if body:
                            body_text = body.decode('utf-8', errors='ignore')
                            
                            metadata = ChunkMetadata(
                                document_id=document_id,
                                section="email_body",
                                chunk_type="body",
                                source_file=file_path,
                                extraction_method="email_parser"
                            )
                            
                            chunks.append(DocumentChunk(
                                content=body_text.strip(),
                                metadata=metadata
                            ))
            else:
                body = msg.get_payload(decode=True)
                if body:
                    body_text = body.decode('utf-8', errors='ignore')
                    
                    metadata = ChunkMetadata(
                        document_id=document_id,
                        section="email_body",
                        chunk_type="body",
                        source_file=file_path,
                        extraction_method="email_parser"
                    )
                    
                    chunks.append(DocumentChunk(
                        content=body_text.strip(),
                        metadata=metadata
                    ))
            
            logger.info(f"Extracted {len(chunks)} elements from email")
            
        except Exception as e:
            logger.error(f"Error parsing email {file_path}: {e}")
            raise
        
        return chunks
    
    def _parse_text(self, file_path: str, document_id: str) -> List[DocumentChunk]:
        """
        Parse plain text files with paragraph-based chunking
        
        Args:
            file_path: Path to text file
            document_id: Unique document identifier
            
        Returns:
            List of DocumentChunk objects
        """
        logger.info(f"Parsing text file: {file_path}")
        chunks = []
        
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            
            # Split by double newlines to preserve paragraph structure
            paragraphs = re.split(r'\n\s*\n', content)
            
            for para_idx, paragraph in enumerate(paragraphs):
                if paragraph.strip():
                    metadata = ChunkMetadata(
                        document_id=document_id,
                        section=f"paragraph_{para_idx + 1}",
                        chunk_type="paragraph",
                        source_file=file_path,
                        extraction_method="text_paragraph_split"
                    )
                    
                    chunks.append(DocumentChunk(
                        content=paragraph.strip(),
                        metadata=metadata
                    ))
            
            logger.info(f"Extracted {len(chunks)} paragraphs from text file")
            
        except Exception as e:
            logger.error(f"Error parsing text file {file_path}: {e}")
            raise
        
        return chunks
    
    def _refine_chunks_semantically(self, chunks: List[DocumentChunk]) -> List[DocumentChunk]:
        """
        Apply semantic refinement to avoid mid-sentence splits using NLTK/spaCy
        
        Args:
            chunks: List of raw document chunks
            
        Returns:
            List of semantically refined chunks
        """
        logger.info("Applying semantic refinement to avoid mid-sentence splits")
        refined_chunks = []
        
        for chunk in chunks:
            try:
                # Use spaCy for better sentence boundary detection if available
                if self.nlp:
                    doc = self.nlp(chunk.content)
                    sentences = [sent.text.strip() for sent in doc.sents if sent.text.strip()]
                else:
                    # Fallback to NLTK sentence tokenization
                    sentences = nltk.sent_tokenize(chunk.content)
                
                # Group sentences into coherent chunks (avoid very small chunks)
                current_chunk = ""
                for sentence in sentences:
                    # If adding this sentence would make chunk too long, create new chunk
                    if len(current_chunk) + len(sentence) > 1000 and current_chunk:
                        # Create chunk with current content
                        refined_metadata = ChunkMetadata(
                            document_id=chunk.metadata.document_id,
                            page_number=chunk.metadata.page_number,
                            section=chunk.metadata.section,
                            chunk_type=chunk.metadata.chunk_type,
                            confidence=0.9,  # Slightly lower confidence for refined chunks
                            source_file=chunk.metadata.source_file,
                            extraction_method=f"{chunk.metadata.extraction_method}_semantic_refined"
                        )
                        
                        refined_chunks.append(DocumentChunk(
                            content=current_chunk.strip(),
                            metadata=refined_metadata
                        ))
                        current_chunk = sentence
                    else:
                        current_chunk += " " + sentence if current_chunk else sentence
                
                # Add remaining content
                if current_chunk.strip():
                    refined_metadata = ChunkMetadata(
                        document_id=chunk.metadata.document_id,
                        page_number=chunk.metadata.page_number,
                        section=chunk.metadata.section,
                        chunk_type=chunk.metadata.chunk_type,
                        confidence=0.9,
                        source_file=chunk.metadata.source_file,
                        extraction_method=f"{chunk.metadata.extraction_method}_semantic_refined"
                    )
                    
                    refined_chunks.append(DocumentChunk(
                        content=current_chunk.strip(),
                        metadata=refined_metadata
                    ))
                    
            except Exception as e:
                logger.warning(f"Error in semantic refinement for chunk: {e}")
                # Keep original chunk if refinement fails
                refined_chunks.append(chunk)
        
        logger.info(f"Semantic refinement complete: {len(chunks)} -> {len(refined_chunks)} chunks")
        return refined_chunks
    
    def _classify_pdf_block(self, block: Dict, text: str) -> str:
        """Classify PDF block type based on formatting and content"""
        # Simple heuristics for block classification
        if len(text) < 50 and text.isupper():
            return "header"
        elif len(text) < 100 and any(keyword in text.lower() for keyword in ['clause', 'section', 'article']):
            return "clause"
        elif 'table' in text.lower() or '\t' in text:
            return "table"
        else:
            return "paragraph"
    
    def _classify_docx_paragraph(self, paragraph) -> str:
        """Classify DOCX paragraph type based on style"""
        style_name = paragraph.style.name.lower() if paragraph.style else ""
        
        if 'heading' in style_name:
            return "header"
        elif 'title' in style_name:
            return "title"
        elif len(paragraph.text) < 100 and any(keyword in paragraph.text.lower() 
                                              for keyword in ['clause', 'section', 'article']):
            return "clause"
        else:
            return "paragraph"
    
    def _extract_table_content(self, table) -> str:
        """Extract structured content from DOCX table"""
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text.strip())
            table_data.append(" | ".join(row_data))
        
        return "\n".join(table_data)
    
    def _is_url(self, path: str) -> bool:
        """Check if path is a URL"""
        return path.startswith(('http://', 'https://'))
    
    def _generate_document_id(self, file_path: str) -> str:
        """Generate unique document ID from file path"""
        return Path(file_path).stem + "_" + str(hash(file_path))[:8]


# Convenience functions for direct usage
def parse_document(file_path_or_url: str) -> List[DocumentChunk]:
    """
    Convenience function to parse a document
    
    Args:
        file_path_or_url: Path to document or URL
        
    Returns:
        List of DocumentChunk objects
    """
    parser = DocumentParser()
    return parser.parse_document(file_path_or_url)


def parse_multiple_documents(file_paths_or_urls: List[str]) -> Dict[str, List[DocumentChunk]]:
    """
    Parse multiple documents efficiently
    
    Args:
        file_paths_or_urls: List of document paths or URLs
        
    Returns:
        Dictionary mapping document paths to chunk lists
    """
    parser = DocumentParser()
    results = {}
    
    for path in file_paths_or_urls:
        try:
            results[path] = parser.parse_document(path)
        except Exception as e:
            logger.error(f"Failed to parse {path}: {e}")
            results[path] = []
    
    return results


if __name__ == "__main__":
    # Example usage and testing
    parser = DocumentParser()
    
    # Test with a sample text
    test_content = """
    This is a test document for the LLM Document Query System.
    
    It contains multiple paragraphs to test the parsing functionality.
    The parser should preserve structure and avoid mid-sentence splits.
    
    This demonstrates the structure-aware chunking approach recommended
    in the methodology document.
    """
    
    # Create temporary test file
    with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False) as f:
        f.write(test_content)
        test_file = f.name
    
    try:
        chunks = parser.parse_document(test_file)
        print(f"Parsed {len(chunks)} chunks:")
        for i, chunk in enumerate(chunks):
            print(f"Chunk {i+1}: {chunk.content[:100]}...")
            print(f"Metadata: {chunk.metadata}")
            print("-" * 50)
    finally:
        os.unlink(test_file)